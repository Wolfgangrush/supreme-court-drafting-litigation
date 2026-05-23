"""Build a proper Bombay-HC-Nagpur reference.docx with locked Word styles.

Pandoc reads styles from this file and applies them to draft-v1.md output.

Layout spec (per bench-config bombay-hc-nagpur.md Section 9):
- A4, Times New Roman 14pt body, 1.5 line spacing
- Margins: 4cm left, 2.5cm right/top/bottom
- Heading 1: TNR 14pt BOLD CENTERED (used for court header, case-number line, INDEX, SYNOPSIS, LIST OF ANNEXURES)
- Heading 2: TNR 14pt BOLD CENTERED with letter-spacing (used for F A C T S, G R O U N D S, P R A Y E R)
- Heading 3: TNR 14pt BOLD LEFT (used for sub-sections inside grounds)
- Body: TNR 14pt justified, 1.5 spacing, 0.5cm first-line indent
- Table: first row bold; explicit cell margins; no auto-fit

Output: bombay-hc-nagpur-reference.docx in same directory as this script.
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("bombay-hc-nagpur-reference.docx")


def set_cell_border(cell, sides=("top", "left", "bottom", "right"), sz=4, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def lock_style(style, *, font="Times New Roman", size_pt=14, bold=False, align=None,
               line_spacing=1.5, color="000000", letter_spacing_pt=None,
               space_before_pt=0, space_after_pt=0, first_line_indent_cm=None,
               keep_with_next=False, outline_level=None):
    font_obj = style.font
    font_obj.name = font
    font_obj.size = Pt(size_pt)
    font_obj.bold = bold
    if color:
        font_obj.color.rgb = RGBColor.from_string(color)
    # Force east-asian font name (some Word installs require this)
    rpr = style.element.get_or_add_rPr()
    for tag in ("rFonts",):
        existing = rpr.find(qn(f"w:{tag}"))
        if existing is not None:
            rpr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rpr.append(rFonts)

    if letter_spacing_pt is not None:
        spacing_el = OxmlElement("w:spacing")
        spacing_el.set(qn("w:val"), str(int(letter_spacing_pt * 20)))
        rpr.append(spacing_el)

    pf = style.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    pf.keep_with_next = keep_with_next

    # Outline level for heading discovery
    if outline_level is not None:
        ppr = style.element.get_or_add_pPr()
        existing = ppr.find(qn("w:outlineLvl"))
        if existing is not None:
            ppr.remove(existing)
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), str(outline_level))
        ppr.append(ol)


def main():
    doc = Document()

    # ----- Page setup: A4, margins per bench-config -----
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ----- Normal (body) style -----
    normal = doc.styles["Normal"]
    lock_style(normal, font="Times New Roman", size_pt=14, bold=False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
               first_line_indent_cm=0.5, space_after_pt=6)

    # ----- Heading 1: court header, INDEX, SYNOPSIS, LIST OF ANNEXURES -----
    h1 = doc.styles["Heading 1"]
    lock_style(h1, font="Times New Roman", size_pt=14, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, color="000000",
               space_before_pt=12, space_after_pt=12, keep_with_next=True,
               outline_level=0)
    # Reset first_line_indent explicitly to 0 for headings
    h1.paragraph_format.first_line_indent = Cm(0)

    # ----- Heading 2: F A C T S / G R O U N D S / P R A Y E R -----
    h2 = doc.styles["Heading 2"]
    lock_style(h2, font="Times New Roman", size_pt=14, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, color="000000",
               letter_spacing_pt=4,  # spaced effect
               space_before_pt=18, space_after_pt=6, keep_with_next=True,
               outline_level=1)
    h2.paragraph_format.first_line_indent = Cm(0)

    # ----- Heading 3: ground sub-headers, prayer sub-clauses, application titles -----
    h3 = doc.styles["Heading 3"]
    lock_style(h3, font="Times New Roman", size_pt=14, bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, color="000000",
               space_before_pt=12, space_after_pt=6, keep_with_next=True,
               outline_level=2)
    h3.paragraph_format.first_line_indent = Cm(0)

    # ----- Title style: for cause-title block (no italic, plain bold centered) -----
    title = doc.styles["Title"]
    lock_style(title, font="Times New Roman", size_pt=14, bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, color="000000",
               space_before_pt=0, space_after_pt=6, keep_with_next=True)
    title.paragraph_format.first_line_indent = Cm(0)

    # ----- BodyText (used after Heading 1 by pandoc) -----
    if "Body Text" in [s.name for s in doc.styles]:
        bt = doc.styles["Body Text"]
        lock_style(bt, font="Times New Roman", size_pt=14, bold=False,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                   first_line_indent_cm=0.5, space_after_pt=6)

    # ----- Compact (list paragraph) -----
    try:
        lp = doc.styles["List Paragraph"]
        lock_style(lp, font="Times New Roman", size_pt=14, bold=False,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                   first_line_indent_cm=0, space_after_pt=6)
    except KeyError:
        pass

    # ----- Demo content: shows the styles at work; pandoc replaces all content with the actual draft -----
    # This is a STYLE TEMPLATE — content here is throwaway. Pandoc reads styles only.
    doc.add_paragraph("IN THE HIGH COURT OF JUDICATURE AT BOMBAY BENCH AT NAGPUR.",
                      style="Heading 1")
    doc.add_paragraph("WRIT PETITION NO. _______ OF 2026", style="Heading 1")
    doc.add_paragraph("(Style template — pandoc replaces this content. Body paragraphs render in Times New Roman 14pt, 1.5 line spacing, justified, with 0.5cm first-line indent.)")
    doc.add_paragraph("F A C T S", style="Heading 2")
    doc.add_paragraph("This is body text inside the F A C T S section. Pandoc applies this style automatically when the markdown source uses a Heading 2 (`## F A C T S`).")
    doc.add_paragraph("G R O U N D S", style="Heading 2")
    doc.add_paragraph("Ground I — opening", style="Heading 3")
    doc.add_paragraph("This is body text inside a numbered ground. Pandoc applies this style automatically.")

    # ----- Table style locked: first row bold centered, fixed col widths -----
    # Show what a table with locked colwidths looks like
    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    # Lock layout (no autofit)
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # Column widths: Sr.No 8% | Annx 8% | Particulars 62% | Date 12% | Pgs 10%
    # Total useable width = A4 21cm - 4cm left - 2.5cm right = 14.5cm
    widths_cm = [1.16, 1.16, 8.99, 1.74, 1.45]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
    # Header row
    hdr = table.rows[0].cells
    headers = ["Sr.No", "Annx", "Particulars", "Date", "Pgs"]
    for idx, cell in enumerate(hdr):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(headers[idx])
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Sample data rows
    data = [
        ["1.", "", "Synopsis", "", ""],
        ["2.", "A", "Sample annexure row", "20.10.2025", ""],
    ]
    for r_idx, row_data in enumerate(data, start=1):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Page numbers — centered at bottom
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    doc.save(str(OUT_PATH))
    print(f"OK · wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
